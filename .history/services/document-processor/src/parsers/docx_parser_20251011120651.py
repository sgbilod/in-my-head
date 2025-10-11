"""
DOCX Parser
Handles Microsoft Word .docx files
"""

from pathlib import Path
import time

from .base_parser import BaseParser, ParsedDocument, ParsingError

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocxParser(BaseParser):
    """Parser for Microsoft Word documents"""

    @property
    def supported_formats(self) -> list[str]:
        return ['.docx']

    async def parse(self, file_path: Path) -> ParsedDocument:
        """
        Parse DOCX file

        Args:
            file_path: Path to the DOCX file

        Returns:
            ParsedDocument with extracted content
        """
        if not DOCX_AVAILABLE:
            raise ParsingError(
                "python-docx library not installed",
                parser=self.name,
                file_path=str(file_path)
            )

        start_time = time.time()

        try:
            doc = Document(file_path)

            # Extract text from paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            text = "\n".join(text_parts)
            text = self._clean_text(text)

            # Extract metadata
            metadata = await self.extract_metadata(file_path)
            core_props = doc.core_properties

            metadata.update({
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'created_at': core_props.created.isoformat()
                if core_props.created else '',
                'modified_at': core_props.modified.isoformat()
                if core_props.modified else '',
                'paragraph_count': len(doc.paragraphs),
                'word_count': len(text.split()),
            })

            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)

            # Extract links
            links = self._extract_links(text)

            parsing_time = time.time() - start_time

            return ParsedDocument(
                text=text,
                text_length=len(text),
                metadata=metadata,
                tables=tables if tables else None,
                links=links if links else None,
                parser_used=self.name,
                parsing_time=parsing_time
            )

        except Exception as e:
            raise ParsingError(
                f"Failed to parse DOCX file: {str(e)}",
                parser=self.name,
                file_path=str(file_path)
            )
