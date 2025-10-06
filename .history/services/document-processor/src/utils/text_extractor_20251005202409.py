"""
Text Extraction Utility
Extracts text content from various document formats
"""

from pathlib import Path
from typing import Optional
import os


async def extract_text(file_path: str, doc_type: str) -> Optional[str]:
    """
    Extract text from a document
    
    Args:
        file_path: Path to document file
        doc_type: Document type (pdf, docx, txt, md)
    
    Returns:
        Extracted text content
    """
    file_path = Path(file_path)
    
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    # Route to appropriate extractor
    if doc_type == 'txt' or doc_type == 'md':
        return await extract_text_file(file_path)
    elif doc_type == 'pdf':
        return await extract_pdf(file_path)
    elif doc_type == 'docx':
        return await extract_docx(file_path)
    else:
        # Try to read as text file
        try:
            return await extract_text_file(file_path)
        except:
            return None


async def extract_text_file(file_path: Path) -> str:
    """
    Extract text from plain text files
    
    Args:
        file_path: Path to text file
    
    Returns:
        File content as string
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        # Try with different encoding
        with open(file_path, 'r', encoding='latin-1') as f:
            return f.read()


async def extract_pdf(file_path: Path) -> str:
    """
    Extract text from PDF files
    
    Args:
        file_path: Path to PDF file
    
    Returns:
        Extracted text content
    """
    try:
        import PyPDF2
        
        text = []
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text.append(page.extract_text())
        
        return '\n\n'.join(text)
    
    except ImportError:
        raise ImportError(
            "PyPDF2 is required for PDF extraction. "
            "Install it with: pip install PyPDF2"
        )
    except Exception as e:
        raise Exception(f"Failed to extract PDF: {str(e)}")


async def extract_docx(file_path: Path) -> str:
    """
    Extract text from DOCX files
    
    Args:
        file_path: Path to DOCX file
    
    Returns:
        Extracted text content
    """
    try:
        import docx
        
        doc = docx.Document(file_path)
        text = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text.append(' | '.join(row_text))
        
        return '\n\n'.join(text)
    
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX extraction. "
            "Install it with: pip install python-docx"
        )
    except Exception as e:
        raise Exception(f"Failed to extract DOCX: {str(e)}")


def get_page_count(file_path: str, doc_type: str) -> Optional[int]:
    """
    Get page count for a document
    
    Args:
        file_path: Path to document
        doc_type: Document type
    
    Returns:
        Number of pages or None
    """
    file_path = Path(file_path)
    
    if doc_type == 'pdf':
        try:
            import PyPDF2
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                return len(pdf_reader.pages)
        except:
            return None
    
    # For other types, estimate from content
    return None
